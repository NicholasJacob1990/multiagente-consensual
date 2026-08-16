#!/usr/bin/env python3
"""Build the editorial DOCX manual from its canonical Markdown source."""

from __future__ import annotations

import argparse
import math
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = RGBColor(31, 41, 55)
ACCENT = RGBColor(30, 64, 175)
MUTED = RGBColor(75, 85, 99)
WHITE = RGBColor(255, 255, 255)
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, name: str, size: float, *, bold: bool | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    fonts = run._element.get_or_add_rPr().rFonts
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), name)


def style_font(style, name: str, size: float, color: RGBColor = INK) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), name)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str, border: str = "CBD5E1") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shading = ppr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        ppr.append(shading)
    shading.set(qn("w:fill"), fill)
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), border)


def force_paragraph_font(paragraph, name: str, size: float) -> None:
    """Apply the editorial font even inside hyperlinks and character styles."""
    half_points = str(int(round(size * 2)))
    for run_element in paragraph._p.iter(qn("w:r")):
        rpr = run_element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run_element.insert(0, rpr)
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            fonts.set(qn(key), name)
        for tag in ("w:sz", "w:szCs"):
            element = rpr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                rpr.append(element)
            element.set(qn("w:val"), half_points)


def set_table_geometry(table, widths: list[int]) -> None:
    """Set redundant OOXML geometry for stable Word and LibreOffice rendering."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, attributes in (
        ("w:tblW", {"w:w": str(TABLE_WIDTH_DXA), "w:type": "dxa"}),
        ("w:tblInd", {"w:w": str(TABLE_INDENT_DXA), "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        element = tbl_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tbl_pr.append(element)
        for key, value in attributes.items():
            element.set(qn(key), value)

    cell_margins = tbl_pr.find(qn("w:tblCellMar"))
    if cell_margins is None:
        cell_margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_margins)
    for side, width in (("top", 70), ("left", 95), ("bottom", 70), ("right", 95)):
        margin = cell_margins.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            cell_margins.append(margin)
        margin.set(qn("w:w"), str(width))
        margin.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def content_aware_widths(table) -> list[int]:
    column_count = max(len(table.columns), 1)
    if column_count == 1:
        return [TABLE_WIDTH_DXA]

    scores: list[float] = []
    for column_index in range(column_count):
        lengths = []
        for row in table.rows:
            text = row.cells[column_index].text.replace("\n", " ").strip()
            lengths.append(min(max(len(text), 4), 80))
        average = sum(lengths) / max(len(lengths), 1)
        peak = max(lengths, default=4)
        scores.append(math.sqrt(average * peak))

    if column_count == 2:
        ratio = scores[0] / max(sum(scores), 1)
        first = int(TABLE_WIDTH_DXA * min(max(ratio, 0.29), 0.48))
        return [first, TABLE_WIDTH_DXA - first]

    minimum = 900 if column_count >= 5 else 1100
    available = TABLE_WIDTH_DXA - minimum * column_count
    total_score = max(sum(scores), 1)
    widths = [minimum + int(available * score / total_score) for score in scores]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is not None:
        return
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def configure_styles(document: Document) -> None:
    styles = document.styles
    for name in ("Normal", "Body Text"):
        if name not in styles:
            continue
        style = styles[name]
        style_font(style, "Lato", 12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.first_line_indent = Inches(0.35)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(5)

    heading_tokens = {
        "Title": (28, ACCENT, 72, 12, WD_ALIGN_PARAGRAPH.CENTER),
        "Subtitle": (15, MUTED, 4, 16, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 1": (20, ACCENT, 18, 10, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 2": (17, ACCENT, 16, 8, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (14.5, MUTED, 13, 6, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 4": (12.5, MUTED, 10, 5, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (size, color, before, after, alignment) in heading_tokens.items():
        if name not in styles:
            continue
        style = styles[name]
        style_font(style, "Lato", size, color)
        style.font.bold = name != "Subtitle"
        fmt = style.paragraph_format
        fmt.alignment = alignment
        fmt.first_line_indent = Inches(0)
        fmt.line_spacing = 1.15
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2", "List Number 3"):
        if name not in styles:
            continue
        style = styles[name]
        style_font(style, "Lato", 12)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(3)

    for name in ("Caption", "Quote", "Intense Quote"):
        if name not in styles:
            continue
        style_font(styles[name], "Lato", 12, MUTED)
        styles[name].paragraph_format.first_line_indent = Inches(0)
        styles[name].paragraph_format.line_spacing = 1.25
        styles[name].paragraph_format.space_after = Pt(6)
    if "Caption" in styles:
        styles["Caption"].font.italic = True
        styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name in ("Source Code", "Code Block"):
        if name not in styles:
            continue
        style = styles[name]
        style_font(style, "Lato", 12, INK)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Inches(0)
        fmt.left_indent = Inches(0.18)
        fmt.right_indent = Inches(0.08)
        fmt.line_spacing = 1.2
        fmt.space_before = Pt(5)
        fmt.space_after = Pt(8)

    for name in ("Default Paragraph Font", "Hyperlink", "Verbatim Char"):
        if name in styles:
            style_font(styles[name], "Lato", 12, ACCENT if name == "Hyperlink" else INK)


def normalize_tables(document: Document) -> None:
    """Rebuild Pandoc tables so Word and LibreOffice share the same layout."""
    for original in list(document.tables):
        matrix = [[cell.text for cell in row.cells] for row in original.rows]
        if not matrix:
            continue
        column_count = max(len(row) for row in matrix)
        replacement = document.add_table(rows=len(matrix), cols=column_count)
        for row_index, values in enumerate(matrix):
            for column_index, value in enumerate(values):
                replacement.cell(row_index, column_index).text = value
        original._tbl.addprevious(replacement._tbl)
        original._element.getparent().remove(original._element)


def configure_document(document: Document) -> None:
    configure_styles(document)
    normalize_tables(document)
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.45)
        section.footer_distance = Inches(0.45)
        section.different_first_page_header_footer = True
        section.first_page_header.paragraphs[0].clear()
        section.first_page_footer.paragraphs[0].clear()

    in_summary = False
    previous_was_heading = False
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""

        if paragraph_index == 1 and text.startswith("Guia prático") and "Subtitle" in document.styles:
            paragraph.style = document.styles["Subtitle"]
            style_name = "Subtitle"
        if text.startswith("Edição 16"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_after = Pt(10)
            force_paragraph_font(paragraph, "Lato", 12)
            previous_was_heading = False
            continue
        if text.startswith("Figura ") and "Caption" in document.styles:
            paragraph.style = document.styles["Caption"]
            style_name = "Caption"

        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        portrait_drawing = False
        if has_drawing:
            for extent in paragraph._p.xpath(".//wp:extent"):
                width = int(extent.get("cx", "1"))
                height = int(extent.get("cy", "1"))
                portrait_drawing = portrait_drawing or height / max(width, 1) > 1.05

        if style_name == "Heading 2":
            paragraph.paragraph_format.page_break_before = True
            in_summary = text == "Sumário"

        if has_drawing:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True
            if portrait_drawing:
                paragraph.paragraph_format.page_break_before = True
            previous_was_heading = False
            continue
        if style_name.startswith("Heading") or style_name in {"Title", "Subtitle", "Caption", "Quote", "Intense Quote", "Block Text"}:
            paragraph.paragraph_format.first_line_indent = Inches(0)
            if style_name == "Caption":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.keep_together = True
                force_paragraph_font(paragraph, "Lato", 12)
            previous_was_heading = style_name.startswith("Heading")
        elif style_name.startswith("List") or style_name in {"Source Code", "Code Block"}:
            paragraph.paragraph_format.first_line_indent = None
            if style_name in {"Source Code", "Code Block"}:
                shade_paragraph(paragraph, "F5F7FA", "1E40AF")
                force_paragraph_font(paragraph, "Lato", 12)
            previous_was_heading = False
        elif in_summary:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.line_spacing = 1.3
            paragraph.paragraph_format.space_after = Pt(3)
            force_paragraph_font(paragraph, "Lato", 12)
            previous_was_heading = False
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Inches(0) if previous_was_heading else Inches(0.35)
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.space_after = Pt(5)
            force_paragraph_font(paragraph, "Lato", 12)
            previous_was_heading = False
        for run in paragraph.runs:
            if style_name in {"Source Code", "Code Block"}:
                rpr = run._r.get_or_add_rPr()
                run_style = rpr.find(qn("w:rStyle"))
                if run_style is not None:
                    rpr.remove(run_style)
                set_font(run, "Lato", 12)
            elif run.font.name is None:
                set_font(run, "Lato", 12)

    for table in document.tables:
        table.style = "Table Grid"
        widths = content_aware_widths(table)
        set_table_geometry(table, widths)
        if table.rows:
            repeat_table_header(table.rows[0])
            for cell in table.rows[0].cells:
                shade_cell(cell, "1E40AF")
        for row_index, row in enumerate(table.rows):
            if row_index > 0 and row_index % 2 == 0:
                for cell in row.cells:
                    shade_cell(cell, "F6F8FB")
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = Inches(0)
                    paragraph.paragraph_format.line_spacing = 1.12
                    paragraph.paragraph_format.space_after = Pt(2)
                    force_paragraph_font(paragraph, "Lato", 12)
                    for run in paragraph.runs:
                        set_font(run, "Lato", 12, bold=True if row_index == 0 else None)
                        if row_index == 0:
                            run.font.color.rgb = WHITE

    max_width = Inches(6.25)
    for shape in document.inline_shapes:
        portrait = shape.height / max(shape.width, 1) > 1.05
        max_height = Inches(7.35 if portrait else 7.65)
        scale = min(max_width / shape.width, max_height / shape.height, 1.0)
        if scale < 1.0:
            shape.width = int(shape.width * scale)
            shape.height = int(shape.height * scale)

    # Keep the first section explicit and stable when Pandoc imports a reference file.
    if document.sections and document.sections[0].start_type != WD_SECTION.NEW_PAGE:
        document.sections[0].start_type = WD_SECTION.NEW_PAGE


def append_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    command = OxmlElement("w:instrText")
    command.set(qn("xml:space"), "preserve")
    command.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, command, separate, end))


def create_clean_reference(path: Path) -> None:
    document = Document()
    configure_document(document)
    for section in document.sections:
        section.different_first_page_header_footer = True
        section.first_page_header.paragraphs[0].clear()
        section.first_page_footer.paragraphs[0].clear()
        header = section.header.paragraphs[0]
        header.text = "MANUAL MULTIAGENTE  •  REFERÊNCIA PRÁTICA"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header.runs:
            set_font(run, "Lato", 9.5, bold=True)
            run.font.color.rgb = MUTED

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prefix = footer.add_run("Edição 16  •  Agosto de 2026  •  Página ")
        set_font(prefix, "Lato", 9.5)
        prefix.font.color.rgb = MUTED
        append_field(footer, "PAGE")
    document.save(path)


def build(source: Path, reference: Path, output: Path, pandoc: str) -> None:
    source = source.resolve()
    reference = reference.resolve()
    output = output.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="manual-editorial-") as raw:
        temporary = Path(raw) / "manual-pandoc.docx"
        clean_reference = Path(raw) / "reference.docx"
        create_clean_reference(clean_reference)
        subprocess.run(
            [
                pandoc,
                str(source),
                "--from=gfm",
                "--to=docx",
                f"--reference-doc={clean_reference}",
                f"--resource-path={source.parent}",
                "--metadata=title:Manual completo — Multiagente Consensual",
                f"--output={temporary}",
            ],
            check=True,
            cwd=source.parent,
        )
        document = Document(temporary)
        configure_document(document)
        staged = Path(raw) / output.name
        document.save(staged)
        os.replace(staged, output)


def export_pdf(docx_path: Path, pdf_output: Path) -> None:
    """Export with the user's font catalog while isolating the LO profile."""
    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice/soffice não encontrado para exportar o PDF.")
    docx_path = docx_path.resolve()
    pdf_output = pdf_output.resolve()
    pdf_output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="manual-pdf-export-") as raw:
        root = Path(raw)
        profile = root / "profile"
        converted = root / "converted"
        profile.mkdir()
        converted.mkdir()
        env = os.environ.copy()
        if sys.platform == "darwin" and Path("/private/tmp").is_dir():
            env["TMPDIR"] = "/private/tmp"
            env["TMP"] = "/private/tmp"
            env["TEMP"] = "/private/tmp"
        subprocess.run(
            [
                soffice,
                f"-env:UserInstallation={profile.resolve().as_uri()}",
                "--invisible",
                "--headless",
                "--norestore",
                "--convert-to",
                "pdf",
                "--outdir",
                str(converted),
                str(docx_path),
            ],
            check=True,
            env=env,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        generated = converted / f"{docx_path.stem}.pdf"
        if not generated.is_file() or generated.stat().st_size == 0:
            raise RuntimeError("LibreOffice não produziu um PDF válido.")
        staged = root / pdf_output.name
        shutil.copy2(generated, staged)
        os.replace(staged, pdf_output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--reference", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--pdf-output", type=Path)
    parser.add_argument("--pandoc", default=shutil.which("pandoc") or "pandoc")
    args = parser.parse_args()
    build(args.source, args.reference, args.output, args.pandoc)
    if args.pdf_output:
        export_pdf(args.output, args.pdf_output)


if __name__ == "__main__":
    main()
